import io
from typing import Any, Dict, List

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _apply_document_setup(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "Heidi Student Workbook | Grade 3\tClassic Books for All"
    header.style = doc.styles["Normal"]

    footer = section.footer.paragraphs[0]
    footer.text = "Lesson 1 | Chapter 1\t"
    footer.add_run("PAGE")


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(146, 64, 14)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(14)


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text or "")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def _shaded_instruction_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EDEDED")
    tc_pr.append(shd)


def _blank_lines(doc: Document, count: int = 3) -> None:
    for _ in range(count):
        doc.add_paragraph("____________________________________________")


def _words_to_know(doc: Document, parsed: Dict[str, Any]) -> None:
    _h2(doc, "Words to Know")
    _body(doc, "Use the book quote to help infer each word's meaning.")

    table = doc.add_table(rows=11, cols=4)
    headers = ["Word", "Definition", "Sentence from book (with page #)", "My own sentence"]
    for c, label in enumerate(headers):
        table.cell(0, c).text = label

    widths = [1.0, 1.8, 2.5, 1.8]
    for col_index, width in enumerate(widths):
        for row in table.rows:
            row.cells[col_index].width = Inches(width)

    vocab = parsed.get("vocabulary", [])[:10]
    for idx in range(10):
        row = idx + 1
        entry = vocab[idx] if idx < len(vocab) else {}
        word_cell = table.cell(row, 0)
        word_cell.text = str(entry.get("WORD", ""))
        for paragraph in word_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        table.cell(row, 1).text = ""
        page_quote = str(entry.get("BOOK_QUOTE", "")).strip()
        table.cell(row, 2).text = page_quote
        table.cell(row, 3).text = ""


def build_docx(parsed_workbook: Dict[str, Any], chapter_json: Dict[str, Any], grade: int) -> io.BytesIO:
    doc = Document()
    _apply_document_setup(doc)

    _h1(doc, "Lesson 1")
    _h2(doc, chapter_json.get("chapter_title", "Going Up to the Alm-Uncle"))

    _h2(doc, "Get Ready to Read")
    _shaded_instruction_box(
        doc,
        str(
            parsed_workbook.get(
                "focus_question",
                "What changes when Heidi goes to her grandfather's home?",
            )
        ),
    )

    _words_to_know(doc, parsed_workbook)

    _h2(doc, "Think About the Story")
    for q in parsed_workbook.get("basic_comprehension", [])[:6]:
        _body(doc, str(q.get("QUESTION", "")))
        _blank_lines(doc, 2)

    _h2(doc, "Reading Between the Lines")
    for q in parsed_workbook.get("inference_questions", [])[:3]:
        _body(doc, str(q.get("QUESTION", "")))
        _blank_lines(doc, 2)

    _h2(doc, "Dig Deeper")
    for q in parsed_workbook.get("analysis_questions", [])[:3]:
        _body(doc, str(q.get("QUESTION", "")))
        _blank_lines(doc, 2)

    _h2(doc, "Multiple Choice")
    for mc in parsed_workbook.get("multiple_choice", [])[:3]:
        _body(doc, str(mc.get("QUESTION", "")))
        _body(doc, f"A) {mc.get('OPTION_A', '')}")
        _body(doc, f"B) {mc.get('OPTION_B', '')}")
        _body(doc, f"C) {mc.get('OPTION_C', '')}")
        _body(doc, f"D) {mc.get('OPTION_D', '')}")

    _h2(doc, "Evidence from the Story")
    for q in parsed_workbook.get("short_answer", [])[:3]:
        _body(doc, str(q.get("QUESTION", "")))
        _blank_lines(doc, 2)

    _h2(doc, "Creative Response")
    cr = parsed_workbook.get("creative_response", {})
    _shaded_instruction_box(
        doc,
        str(cr.get("PROMPT", "Write a letter to a character from this chapter.")),
    )
    _body(doc, f"• {cr.get('HINT_1', '')}")
    _body(doc, f"• {cr.get('HINT_2', '')}")
    _body(doc, f"• {cr.get('HINT_3', '')}")
    _body(doc, "Dear [character],")
    _blank_lines(doc, 12)
    _body(doc, "Sincerely,")
    _body(doc, "[Student]")

    _h2(doc, "Writing Rubric")
    rubric = doc.add_table(rows=5, cols=5)
    rubric.cell(0, 0).text = "WRITING RUBRIC (16 POINTS POSSIBLE)"
    rubric.cell(1, 0).text = "Ideas"
    rubric.cell(2, 0).text = "Details from the Chapter"
    rubric.cell(3, 0).text = "Writing Clarity"
    rubric.cell(4, 0).text = "Spelling/Grammar"

    _h2(doc, "Character Chart")
    ctable = doc.add_table(rows=6, cols=3)
    ctable.cell(0, 0).text = "Character Name"
    ctable.cell(0, 1).text = "What They Look Like and How They Act"
    ctable.cell(0, 2).text = "What This Shows About Them"
    for i, char in enumerate(parsed_workbook.get("character_chart", [])[:5], start=1):
        ctable.cell(i, 0).text = str(char.get("CHARACTER", ""))
        ctable.cell(i, 1).text = ""
        ctable.cell(i, 2).text = ""

    _h2(doc, "Draw It")
    _body(
        doc,
        str(
            parsed_workbook.get(
                "draw_it_prompt",
                "Draw what you imagine a scene or character looks like. Use the description from the book to guide your drawing.",
            )
        ),
    )
    _blank_lines(doc, 6)

    _h2(doc, "Reflect on Drawing")
    default_stems = [
        "My drawing shows...",
        "I chose these colors/details because...",
        "This part of the story is important because...",
    ]
    stems = parsed_workbook.get("reflect_stems", [])[:3] or default_stems
    for stem in stems:
        _body(doc, str(stem))
        _blank_lines(doc, 1)

    _h2(doc, "Bonus Challenge - Follow the Story")
    _body(doc, "Number these story events in the correct order (1-7).")
    for idx, event in enumerate(parsed_workbook.get("timeline_events", [])[:7], start=1):
        _body(doc, f"{idx}. {event}")

    _h2(doc, "Thinking Deeper")
    _body(
        doc,
        str(
            parsed_workbook.get(
                "thinking_deeper_question",
                parsed_workbook.get("prediction_setup", "What do you think matters most in this chapter, and why?"),
            )
        ),
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
