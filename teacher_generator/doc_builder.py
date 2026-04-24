import io
from typing import Any, Dict

from docx import Document
from docx.shared import Inches, Pt


def _setup(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(14)


def _text(doc: Document, text: str) -> None:
    doc.add_paragraph(text or "")


def build_docx(parsed: Dict[str, Any], chapter_json: Dict[str, Any], grade: int) -> io.BytesIO:
    doc = Document()
    _setup(doc)

    title = doc.add_paragraph()
    tr = title.add_run(f"Teacher Guide — Lesson 1 — {chapter_json.get('chapter_title', '')}")
    tr.bold = True
    tr.font.size = Pt(18)

    _h2(doc, "Get Ready to Read")
    _text(doc, parsed.get("get_ready_to_read", ""))

    _h2(doc, "Words to Know")
    for word in parsed.get("vocabulary", []):
        _text(doc, f"{word.get('WORD', '')} (p.{word.get('PAGE', '')})")
        _text(doc, f"Definition: {word.get('KID_FRIENDLY_DEFINITION', '')}")
        _text(doc, f"Context: {word.get('CONTEXT_SENTENCE', '')}")
        _text(doc, f"Example: {word.get('EXAMPLE_SENTENCE', '')}")

    _h2(doc, "Guided Reading")
    for sec in parsed.get("guided_sections", []):
        _text(doc, f"Section {sec.get('SECTION_NUMBER', '')} ({sec.get('PAGE_RANGE', '')})")
        _text(doc, sec.get("PAUSE_QUESTIONS", ""))
    _text(doc, f"Teaching tip: {parsed.get('guided_tip', '')}")

    _h2(doc, "Think About the Story (Teacher Notes)")
    _text(doc, parsed.get("think_about_answers", ""))

    _h2(doc, "Reading Between the Lines (Teacher Notes)")
    _text(doc, parsed.get("reading_between_answers", ""))

    _h2(doc, "Dig Deeper (Teacher Notes)")
    _text(doc, parsed.get("dig_deeper_answers", ""))

    _h2(doc, "Multiple Choice (Answer Key)")
    _text(doc, parsed.get("multiple_choice_answers", ""))

    _h2(doc, "Evidence from the Story (Answer Key)")
    _text(doc, parsed.get("short_answer_answers", ""))

    _h2(doc, "Creative Response Guide")
    _text(doc, parsed.get("creative_response_guide", ""))

    _h2(doc, "Teaching Notes & Differentiation")
    _text(doc, parsed.get("lesson_overview", ""))
    _text(doc, parsed.get("measurable_objectives", ""))
    _text(doc, parsed.get("standards_list", ""))
    _text(doc, parsed.get("materials_needed", ""))
    _text(doc, f"Struggling Readers: {parsed.get('vocab_diff_struggling', '')}")
    _text(doc, f"ELL: {parsed.get('vocab_diff_ell', '')}")

    _h2(doc, "Character Chart (Answer Key)")
    _text(doc, parsed.get("character_chart_answer_key", ""))

    _h2(doc, "Draw It + Reflection (Answers)")
    _text(doc, parsed.get("exit_ticket_guide", ""))

    _h2(doc, "Bonus Challenge (Answer)")
    _text(doc, parsed.get("differentiation_summary", ""))

    _h2(doc, "Thinking Deeper (Timeline Answer Key)")
    _text(doc, parsed.get("timeline_answer_key", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
